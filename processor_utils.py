import os
import re
import shutil
import fitz  # PyMuPDF
import pdfplumber
import pandas as pd
from docx import Document as DocxDocument
from collections import Counter
from PIL import Image
import io
import urllib.parse
import pytesseract

# --- 💡 满血功能：视觉大模型依赖 ---
try:
    import google.generativeai as genai

    HAS_VISION_AI = True
except ImportError:
    HAS_VISION_AI = False

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

# ==========================================
# ⚙️ 核心配置区
# ==========================================
# 1. OCR 引擎配置 (本地 Windows 时取消注释)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# 2. 视觉大模型配置 (看图识表功能)
VISION_API_KEY = "YOUR_API_KEY_HERE"  # 填入你的 Gemini Key 或置空
ENABLE_VISION_AI = HAS_VISION_AI and VISION_API_KEY != "YOUR_API_KEY_HERE"

if ENABLE_VISION_AI:
    genai.configure(api_key=VISION_API_KEY)
    vision_model = genai.GenerativeModel('gemini-1.5-flash')


def extract_markdown_table_with_ai(image_path):
    """💡 满血功能：让 AI 看着截图，直接手敲出 Markdown 表格"""
    if not ENABLE_VISION_AI:
        return "<IMAGE_ONLY>"
    try:
        img = Image.open(image_path)
        prompt = """
        你是一个极其严谨的学术数据提取专家。
        请仔细观察这张图片，如果里面包含表格或数据图表，请严格将其转换为 Markdown 格式的表格。
        要求：
        1. 完全忠实于原图数据，不可捏造。
        2. 只输出纯粹的 Markdown 表格文本，不要包裹 ``` 代码块。
        3. 如果这完全不是一个表格（比如纯折线图、纯柱状图、设备照片），请严格只回复："<IMAGE_ONLY>"。
        """
        response = vision_model.generate_content([prompt, img])
        result = response.text.strip()

        # 清洗代码块标记
        if result.startswith("```markdown"):
            result = result[11:-3].strip()
        elif result.startswith("```"):
            result = result[3:-3].strip()

        return result
    except Exception as e:
        print(f"视觉解析预警: {e}")
        return "<IMAGE_ONLY>"


def merge_rects(rect_list, x_tol=10, y_tol=15):
    """合并相互重叠或靠近的矩形框"""
    merged = []
    for r in rect_list:
        found = False
        for i, m in enumerate(merged):
            if r.intersects(m + (-x_tol, -y_tol, x_tol, y_tol)):
                merged[i] = m | r
                found = True
                break
        if not found: merged.append(r)
    return merged


def advanced_pdf_parser(file_path, file_name):
    """
    终极学术解析管线：去水印 + 图表拦截 + 公式雷达 + AI视觉重建 + 双栏重排
    """
    pdf_stem = os.path.splitext(file_name)[0]
    asset_base_dir = "extracted_assets"
    pdf_asset_dir = os.path.join(asset_base_dir, pdf_stem)

    # 确保每个PDF的图片拥有独立的专属子文件夹
    if os.path.exists(pdf_asset_dir):
        shutil.rmtree(pdf_asset_dir)
    os.makedirs(pdf_asset_dir, exist_ok=True)

    doc_fitz = fitz.open(file_path)
    doc_plumber = pdfplumber.open(file_path)

    # 1. 跨页天眼：幽灵水印扫描
    coord_counter = Counter()
    for page in doc_fitz:
        items = [fitz.Rect(img["bbox"]) for img in page.get_image_info()]
        items += [fitz.Rect(d["rect"]) for d in page.get_drawings()]
        for r in items:
            fuzzy_coord = (round(r.x0, -1), round(r.y0, -1), round(r.width, -1), round(r.height, -1))
            coord_counter[fuzzy_coord] += 1
    watermark_blacklist = {c for c, count in coord_counter.items() if count >= 2}

    final_md_blocks = []

    # 2. 逐页深度解析核心循环
    for page_num in range(len(doc_fitz)):
        page_f = doc_fitz[page_num]
        page_p = doc_plumber.pages[page_num]
        page_w, page_h = page_f.rect.width, page_f.rect.height
        page_content = [f"## {file_name} - 第 {page_num + 1} 页\n"]

        hard_mask_rects = []
        soft_asset_rects = []
        blocks = page_f.get_text("blocks")

        # --- A. 探测图片区 ---
        for img in page_f.get_image_info():
            r = fitz.Rect(img["bbox"])
            fuzzy_r = (round(r.x0, -1), round(r.y0, -1), round(r.width, -1), round(r.height, -1))
            if r.width < page_w * 0.9 and fuzzy_r not in watermark_blacklist:
                hard_mask_rects.append(r)

        # --- B. 探测有线表格(三线表) ---
        drawings = page_f.get_drawings()
        h_lines = sorted([fitz.Rect(d["rect"]) for d in drawings if
                          fitz.Rect(d["rect"]).width > 40 and fitz.Rect(d["rect"]).height < 5], key=lambda x: x.y0)
        if h_lines:
            temp = [h_lines[0]]
            for line in h_lines[1:]:
                if 0 <= (line.y0 - temp[-1].y1) < 250:
                    temp.append(line)
                else:
                    if len(temp) >= 2:
                        hard_mask_rects.append(fitz.Rect(min(l.x0 for l in temp) - 5, min(l.y0 for l in temp) - 10,
                                                         max(l.x1 for l in temp) + 5, max(l.y1 for l in temp) + 10))
                    temp = [line]
            if len(temp) >= 2:
                hard_mask_rects.append(
                    fitz.Rect(min(l.x0 for l in temp) - 5, min(l.y0 for l in temp) - 10, max(l.x1 for l in temp) + 5,
                              max(l.y1 for l in temp) + 10))

        # --- C. 探测无边框表格 ---
        ts = {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "snap_y_tolerance": 5,
            "min_words_vertical": 3,
            "intersection_x_tolerance": 15
        }
        try:
            for t in page_p.find_tables(table_settings=ts):
                soft_asset_rects.append(fitz.Rect(t.bbox))
        except:
            pass

        # --- 💡 满血功能：数学公式雷达 ---
        math_symbols = ['∑', '∏', '≈', '∝', '∈', '∫', '≥', '≤', '≠', '±', '=', '+', '-', '×', '÷']
        for b in blocks:
            if b[6] == 0:
                text = b[4].strip()
                has_eq_num = bool(re.search(r'[\(（\[]\s*\d+\s*[\)）\]]$', text))
                has_math_op = any(sym in text for sym in math_symbols)
                if (has_eq_num and has_math_op) or text.startswith("式中") or text.startswith("其中，"):
                    # 保护公式区域
                    soft_asset_rects.append(fitz.Rect(b[:4]) + (-5, -5, 5, 5))

        merged_hard_masks = merge_rects(hard_mask_rects)
        merged_soft_assets = merge_rects(soft_asset_rects)
        all_screenshot_assets = merge_rects(merged_hard_masks + merged_soft_assets)

        # --- 📸 截图执行与 AI 介入 ---
        asset_entries = []
        for a_idx, a_rect in enumerate(all_screenshot_assets):
            if a_rect.width < 20 or a_rect.height < 15: continue
            if a_rect.width > page_w * 0.95 and a_rect.height > page_h * 0.8: continue

            asset_filename = f"p{page_num + 1}_a{a_idx}.png"
            save_path = os.path.join(pdf_asset_dir, asset_filename)
            clip_rect = a_rect.intersect(page_f.rect)
            page_f.get_pixmap(clip=clip_rect, matrix=fitz.Matrix(3, 3)).save(save_path)

            safe_stem = urllib.parse.quote(pdf_stem)
            safe_file = urllib.parse.quote(asset_filename)

            # 💡 核心修改：去掉 ../，直接使用相对于项目根目录的路径，便于前端大模型准确吐出并被 UI 拦截
            md_path = f"{asset_base_dir}/{safe_stem}/{safe_file}"

            # 💡 满血功能：AI 视觉重建结构化表格
            ai_markdown = extract_markdown_table_with_ai(save_path)
            if ai_markdown == "<IMAGE_ONLY>":
                # 退化为 OCR 摘要
                ocr_res = ""
                try:
                    img_pil = Image.open(
                        io.BytesIO(page_f.get_pixmap(clip=clip_rect, matrix=fitz.Matrix(2, 2)).tobytes()))
                    ocr_res = pytesseract.image_to_string(img_pil, lang='chi_sim+eng').strip().replace("\n", " ")
                except:
                    pass
                tag = f"\n![图表资产]({md_path})\n> 📌 **[内容摘要]**: {ocr_res if ocr_res else '视觉内容'}\n"
            else:
                tag = f"\n![图表资产]({md_path})\n> 📊 **[AI 重建结构化表格]**\n{ai_markdown}\n"

            asset_entries.append({'rect': a_rect, 'tag': tag})

        # --- 📝 正文提取与 N-Column 多栏自适应 ---
        raw_text_check = "".join([b[4] for b in blocks if b[6] == 0]).strip()

        # 扫描件 OCR 兜底检测
        if len(re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]', raw_text_check)) < 100:
            try:
                pix = page_f.get_pixmap(matrix=fitz.Matrix(2.5, 2.5))
                ocr_full = pytesseract.image_to_string(Image.open(io.BytesIO(pix.tobytes())), lang='chi_sim+eng')
                page_content.append(ocr_full + "\n\n")
            except Exception as e:
                page_content.append(f"⚠️ 扫描件 OCR 解析失败: {e}\n\n")
        else:
            # 过滤掉图表重叠区的文本
            valid_text = [b for b in blocks if
                          b[6] == 0 and not any(fitz.Rect(b[:4]).intersects(m) for m in merged_hard_masks)]

            mid = page_w / 2
            elements = []
            for b in valid_text:
                elements.append({'type': 'text', 'y': b[1], 'x': b[0], 'val': b[4]})
            for e in asset_entries:
                elements.append({'type': 'img', 'y': e['rect'].y0, 'x': e['rect'].x0, 'val': e['tag']})

            # 双栏还原
            left = sorted([e for e in elements if e['x'] < mid], key=lambda e: e['y'])
            right = sorted([e for e in elements if e['x'] >= mid], key=lambda e: e['y'])

            for e in left + right:
                if e['type'] == 'text':
                    clean_text = e['val'].replace("-\n", "").replace("\n", "").strip()
                    if clean_text:
                        # 💡 满血功能：学术引文高亮
                        clean_text = re.sub(r'(\[\d+(?:,\s*\d+|-?\d+)*\])', r' **[引用: \1]** ', clean_text)
                        clean_text = re.sub(r'(\[出处:.*?\])', r' **\1** ', clean_text)
                        page_content.append(clean_text + "\n\n")
                else:
                    page_content.append(e['val'] + "\n\n")

        final_md_blocks.append("".join(page_content))

    doc_fitz.close()
    doc_plumber.close()
    return "\n---\n".join(final_md_blocks)


def extract_text_from_any(file_path, file_name):
    """路由分发：根据不同格式调用不同解析器"""
    ext = os.path.splitext(file_name)[1].lower()
    if ext == '.pdf':
        return advanced_pdf_parser(file_path, file_name)
    elif ext == '.docx':
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    elif ext in ['.xlsx', '.xls', '.csv']:
        # 将表格转化为 Markdown 对齐结构，消除幻觉
        df = pd.read_csv(file_path) if ext == '.csv' else pd.read_excel(file_path)
        return "这是一份表格数据内容：\n\n" + df.to_markdown(index=False)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    return ""


def process_and_ingest(uploaded_file, scope_tag, db_dir="chroma_db"):
    """
    满血版流程：
    1. 保存上传文件到临时目录
    2. 进行万能深度解析
    3. 将解析出的干净 Markdown 保存到 data_source 文件夹 (生成 .md 文件)
    4. 对 Markdown 切分并执行向量化 RAG 入库
    5. 删除临时源文件
    """
    try:
        DATA_SOURCE_DIR = "data_source"
        os.makedirs(DATA_SOURCE_DIR, exist_ok=True)

        # 1. 保存流对象为物理临时文件
        os.makedirs("temp_uploads", exist_ok=True)
        temp_path = os.path.join("temp_uploads", uploaded_file.name)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # 2. 万能深度解析 (PDF走图文管线，表格走 Markdown)
        md_content = extract_text_from_any(temp_path, uploaded_file.name)

        if not md_content or not md_content.strip():
            print("内容解析为空。")
            return False

        # 3. 物理落地：将解析后的干净内容生成为 .md 文件，存入 data_source
        md_filename = os.path.splitext(uploaded_file.name)[0] + ".md"
        md_filepath = os.path.join(DATA_SOURCE_DIR, md_filename)
        with open(md_filepath, "w", encoding="utf-8") as f:
            f.write(md_content)

        # 4. 向量化入库 (喂入生成好的 Markdown 字符串)
        doc_obj = [Document(page_content=md_content, metadata={"source": uploaded_file.name, "scope": scope_tag})]
        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=80)
        splits = splitter.split_documents(doc_obj)

        embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-zh-v1.5")
        Chroma.from_documents(documents=splits, embedding=embeddings, persist_directory=db_dir)

        # 5. 清理临时二进制原文件
        if os.path.exists(temp_path):
            os.remove(temp_path)

        return True
    except Exception as e:
        print(f"处理失败详情: {e}")
        return False